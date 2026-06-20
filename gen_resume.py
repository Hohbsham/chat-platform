"""Generate Chinese resume from English original, highlight Agent development."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

def add_heading_text(text, size=14, bold=True, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    if color:
        run.font.color.rgb = color
    return p

def add_section(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_line(text, indent=False, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5) if indent else Cm(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'Microsoft YaHei'
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Microsoft YaHei'
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("• " + text)
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    return p

# ═══ HEADER ═══
add_heading_text("杨子旭", size=22, color=RGBColor(0, 51, 102), space_after=2)
add_line("Yangzixu2004@outlook.com  |  +86 18717199788  |  github.com/Hohbsham")
add_line("AI Agent 开发工程师 | 多智能体协作系统 | 具身智能 | AIGC")

# ═══ 研究愿景 ═══
add_section("研究愿景")
add_bullet("致力于训练能够深度适配具体项目工作流的领域 AI Agent，使其不仅理解任务，更理解项目的上下文、规范和长期目标")
add_bullet("探索 Agent 间协作的技术路径：通过 A2A (Agent-to-Agent) 协议实现 Agent 间能力发现与任务委派，通过 MCP (Model Context Protocol) 打通工具生态，通过 ACP (Agent Communication Protocol) 标准化 Agent 间对话，让多个专业 Agent 像高效团队一样分工、沟通、互审、迭代")
add_bullet("核心信念：单个 Agent 的能力存在上限，但通过标准化的通信协议和协作机制，多个专业化 Agent 协同可以产生 1+1 远大于 2 的效果——这正是 Agent 协作网络的意义所在")

# ═══ 求职意向 ═══
add_section("求职意向")
add_line("AI Agent 开发工程师 / 多智能体系统工程师 / 大模型应用开发", indent=True)
add_line("2027届本科 · 期望实习/全职 · 可尽快到岗", indent=True)

# ═══ 教育背景 ═══
add_section("教育背景")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.space_before = Pt(0)
r1 = p.add_run("澳门科技大学 (MUST)")
r1.bold = True; r1.font.size = Pt(10.5)
r2 = p.add_run("                                                    澳门, 中国")
r2.font.size = Pt(10.5)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run("人工智能 理学学士")
r1.bold = True; r1.font.size = Pt(10.5)
r2 = p.add_run("                                               2023年9月 - 2027年6月 (预期)")
r2.font.size = Pt(10.5)

add_bullet("文化学术杰出奖 (Merit-based Outstanding Cultural and Academic Award), 2023/24 学年")

# ═══ Agent 开发核心能力 (重点) ═══
add_section("Agent 开发核心能力")

add_line("多智能体通信与编排")
add_bullet("从零构建 A2A (Agent-to-Agent) 通信平台，探索 Agent 协作网络的工程实现：Agent 能力注册 → 任务自动匹配 → 多模式协作（广播/群聊/审批/任务依赖链）")
add_bullet("深入调研 Google A2A、MCP、ACP 三大 Agent 协议栈，以及 CrewAI、AutoGen、LangGraph、OpenClaw 等主流 Agent 框架的架构设计，撰写 5 框架对比分析报告")
add_bullet("实现 WebSocket + HTTP REST 双通道实时通信；Agent 身份模型借鉴 CrewAI (Role/Goal/Backstory)，Task 状态机遵循 A2A Protocol 规范 (pending→claimed→working→input_required→completed)")
add_bullet("集成 DeepSeek v4 Pro API 驱动 Agent 自主回复，实现感知→推理→响应完整闭环；设计 GroupChat 发言人选机制 (AutoGen 参考) 与 Human-in-the-Loop 审批流程 (LangGraph 参考)")
add_bullet("零外部依赖纯 Python stdlib 服务端 + 原生 WebSocket 前端，支持 Windows 服务持久化，已开源在 GitHub")

add_line("")
add_line("具身智能 Agent (Embodied AI)")
add_bullet("LanderPi 项目：设计并实现语音控制自主导航 Agent — ASR (Whisper-1) → LLM (Llama-3.1-8B) 语义理解 → Nav2 自主规划执行，完整感知-推理-行动闭环")
add_bullet("Raspberry Pi 5 + ROS2 Humble 全栈集成，融合 LiDAR、深度相机、麦克风阵列等多模态传感器，实现 RTAB-VSLAM 建图与 AMCL 定位")
add_bullet("解决 ARM 嵌入式平台上 ROS2 集成的关键工程挑战（TF 时间戳同步、AMCL 生命周期管理、FastDDS 共享内存通信等）")

add_line("")
add_line("领域 Agent 应用")
add_bullet("Buffett-Munger Agent：构建投资智囊 AI Agent，应用巴菲特与芒格的价值投资哲学，实现七维分析框架 (护城河/管理层/逆向思考等)")
add_bullet("基于 RL 强化学习 + 人类反馈训练，整合 23 本投资经典 + 181 份股东信的知识库，Agent 行为随反馈永久进化")
add_bullet("设计完整的 Agent 训练闭环：用户提问 → 七维模板分析 → 人工打分反馈 → CLAUDE.md 自动更新 → 行为永久改变")

add_line("")
add_line("大模型与 AI 工具链")
add_bullet("熟练使用 Claude Code、Cursor、OpenClaw Agent 等 AI 编程 Agent 进行日常开发与项目管理")
add_bullet("熟悉 OpenAI / Anthropic / DeepSeek API 调用范式、上下文压缩、Agent 工具使用 (Tool Use) 机制")
add_bullet("了解 MCP (Model Context Protocol) 与 A2A 协议的应用场景与互补关系")

add_line("")
add_line("数据集构建与处理")
add_bullet("具备大规模视觉数据集构建经验：从数据生成、自动标注到质量审核的全流程设计与实现")
add_bullet("设计多 Agent 协作的数据生产架构（Generation Agent + Segmentation Agent + Annotation Agent 协同），支持万级样本的高效产出")
add_bullet("构建 5000+ 样本多层级服装解析数据集，设计 MD5 去重 + Qwen-VL 自动标注 + JSONL 规范化 Pipeline")

add_line("")
add_line("Buffett-Munger Agent — Agent 训练方法论与领域专家系统  |  RLHF, 知识库工程, Claude Code Skills")
add_line("github.com/Hohbsham/buffett-munger-agent", indent=True)
add_bullet("提出\"Markdown文件 + 人类评分 = 专家Agent\"的轻量级 Agent 训练方法论：无需 GPU、无需标注数据、无需写代码，仅通过 CLAUDE.md 迭代 + 人类反馈即可训练出领域专家 Agent")
add_bullet("设计完整训练闭环：Agent 输出分析 → 多维度人工评分 (1-5⭐) → 即时反馈改进 → CLAUDE.md 固化行为 → 下次自动加载训练好的版本，4 轮从泛化到精化")
add_bullet("训练成果：形成稳定七维分析框架 + 🟢🟡🔴 红绿灯总结表 + 竞品对比 + 精确击球点/卖出价；知识库涵盖 181 份巴菲特股东信 (1957-2025) + 23 本价值投资经典")
add_bullet("封装 3 个可复用技能 (/book-tutor, /deep-analysis, /deep-research)，验证\"行为封装\"模式——稳定能力封装为命令，无需每次重复指令")
add_bullet("项目定位为 Agent 训练范本：他人可克隆仓库、替换 CLAUDE.md 中的领域知识、按同样流程训练自己的专家 Agent")

# ═══ 研究项目 (准备发表) ═══
add_section("研究项目 (准备发表)")

add_line("BPT: Garment Mesh Regularization via Point-Cloud Reconstruction", bold_prefix="")
add_line("PyTorch, MeshTransformer, 3D Vision, AutoDL H20  |  导师: 林诚 教授")
add_bullet("研究问题：如何从含噪声的稀疏点云（2048点）中重建出规整、拓扑一致的四边形服装 Mesh，服务于虚拟试衣与数字人建模")
add_bullet("方法：基于 711M 参数 BPT (Blocked-Parallel Transformer) 的端到端重建框架，直接从点云预测离散 Mesh Token 序列")
add_bullet("构建高质量数据集 (2,807 训练 / 311 独立测试)，预计算 3,117 个 Embedding 缓存以加速训练迭代")
add_bullet("在 AutoDL H20-NVLink 96GB GPU 上实现 fp32 全精度训练，经历 20+ 次崩溃修复（NaN 污染、梯度爆炸、dtype 不匹配等），最终形成稳定训练方案")
add_bullet("评估指标：Chamfer Distance (%bbox)、Hausdorff、Face Coverage、Quad Ratio；目标 Chamfer <2%, Quad Ratio >90%")

add_line("")
add_line("Agent-Driven Research: AutoResearch Pipeline", indent=True)
add_bullet("设计闭环 Agent 科研流水线：program.md 定义目标 → Agent 迭代修改训练代码 → SSH 远程部署到 H20 GPU → monitor.py 每 30 分钟自动巡检（崩溃检测/自动重启/NaN 降学习率）→ 训练完成自动推理评估 → 评估结果反馈给 Agent 决策下一轮超参数")
add_bullet("Agent 自主完成 20+ 轮训练迭代，解决 fp16/fp32 混合精度崩溃、NaN 梯度污染权重、.pyc 缓存覆盖新代码等底层工程问题")
add_bullet("实现完整 AutoResearch 闭环：错误分类与自动修复决策树、GPU 适配速查表（6 种 GPU）、训练监控与自动调参逻辑")

add_line("")
add_line("MultiLayer-VTON: Agent-Driven Multi-Layer Garment Parsing Dataset", bold_prefix="")
add_line("Python, SAM, Doubao API, Qwen-VL, Grounding DINO  |  导师: 林诚 教授")
add_bullet("研究问题：现有虚拟试衣数据集缺乏多层级、可独立编辑的服装图层标注，限制了分层编辑与精细化试衣任务的研究")
add_bullet("构建大规模多层级服装解析数据集：从裸模图像出发，利用生成模型 + 分割模型构建包含独立图层的服装数据（单套搭配含 4-6 个独立图层）")
add_bullet("设计 8 步数据生成 Pipeline：裸模生成 → 去背 → 多 Layer 衣服生成 → SAM2 分割 → 笛卡尔积组合 → Mask 生成 → Qwen-VL 自动标注 → JSONL 规范化输出")

add_line("")
add_line("Agent-Driven Data Generation: Multi-Agent Collaborative Pipeline", indent=True)
add_bullet("设计多 Agent 协作的数据生产架构：Generation Agent（豆包 API）负责生成服装图层 → Segmentation Agent（Grounding DINO + SAM2.1）负责精细化分割 → Annotation Agent（Qwen-VL）负责自动标注属性")
add_bullet("不同 Agent 各司其职形成闭环：Generation Agent 和 Segmentation Agent 之间的迭代反馈（分割失败 → 调整生成参数重新生成）；Annotation Agent 独立运行，MD5 去重后标注并扩展到全量样本")
add_bullet("产出两版高质量数据集：Dataset.34（100 裸模 × 4 层 × 笛卡尔积扩展）和 Dataset.35（250 裸模 × 6 层独立 × 多样性体型/肤色/年龄），支持分层虚拟试衣研究")

# ═══ 项目经历 ═══
add_section("项目经历")

add_line("A2A Agent 通信平台  |  Python, WebSocket, DeepSeek API, A2A/MCP/ACP")
add_line("github.com/Hohbsham/chat-platform", indent=True)
add_bullet("轻量级多 Agent 协作平台，实现 Agent 能力注册 → 任务自动匹配 → AI 驱动回复 → 实时聊天 UI 全流程可见")
add_bullet("实现 4 种协作模式：广播任务 (Broadcast)、群聊轮次 (GroupChat, AutoGen)、人工审批 (HITL, LangGraph)、任务依赖链 (Context Chaining, CrewAI)")
add_bullet("Agent 身份模型借鉴 CrewAI (Role/Goal/Backstory)，Task 状态机遵循 A2A Protocol 规范，WebSocket 实时推送方案参考 OpenClaw 架构")
add_bullet("愿景：作为 Agent 协作网络的实验平台，验证 A2A 协议下的能力发现、MCP 工具生态打通、ACP 标准化的 Agent 对话等关键协作机制")

add_line("")
add_line("LanderPi - 具身 AI Agent 机器人  |  Python, ROS2, LLM, SLAM")
add_line("github.com/Hohbsham/LanderPi", indent=True)
add_bullet("语音指令 -> LLM 意图理解 -> 自主导航的完整 Agent 闭环系统")
add_bullet("Raspberry Pi 5 + ROS2 Humble 全栈集成，约 110 个诊断控制脚本")

# ═══ 研究经历 ═══
add_section("研究经历")
add_line("智能计算实验室, 澳门科技大学                                                   澳门, 中国")
add_line("本科科研助理, 导师: 林诚 教授                                        2025年9月 - 至今", indent=True)
add_bullet("研究方向: AIGC (AI-Generated Content), 聚焦生成式模型与内容合成技术")
add_bullet("与研究生合作研究大规模生成式架构，参与文献调研、模型实现与实验评估")

# ═══ 技能 ═══
add_section("技能")
add_bullet("编程语言: Python, C++, R, SQL")
add_bullet("AI/ML: TensorFlow Lite, Edge AI, AIGC, 生成式模型, LLM 集成与 API 调用")
add_bullet("Agent 开发: Claude Code, Cursor, OpenClaw, A2A 系统设计与实现, 多 Agent 编排")
add_bullet("机器人: ROS2 Humble, Navigation2, SLAM (RTAB-Map), AMCL, LiDAR, 传感器融合 (EKF/IMU)")
add_bullet("语言: 中文 (母语), 英语 (流利)")

# ═══ 校园活动 ═══
add_section("校园活动")
add_line("SUMUSTACGNC 动漫社, 澳门科技大学                              2023年11月 - 2025年5月")
add_bullet("参与社团活动策划与漫展摊位设计")

# ── Save ──
output = r"C:\Users\YANGZ\Desktop\杨子旭_简历_Agent开发_v5.docx"
doc.save(output)
print(f"Saved: {output}")
